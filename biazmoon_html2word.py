from docx import Document
from bs4 import BeautifulSoup
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
mystyle = document.styles.add_style('mystyle', WD_STYLE_TYPE.CHARACTER)

with open("VHDL_my_question_list.html", encoding='utf8') as fp: # you should replace your file with mine
    soup = BeautifulSoup(fp, 'html.parser')
    question_and_choices_list  = list(soup.find_all(['p', 'span'])) # find and get all question and choices
    
CHOICE_LIST = [' (الف', ' (ب', ' (ج', ' (د'] # choice list for questions

with open("testfile.txt", 'w') as output_text:
    qn = 1 # question number
    cn = 0 # choice number(index)
    # qn_cached = qn
    for tag in question_and_choices_list:
        if tag.name == 'p': # if it's question
            line = f'\n{qn}- {tag.text.strip()}'
            run = document.add_paragraph() # write to doc file
            run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = run.add_run(line)
            run.style = mystyle
            font = run.font
            font.rtl = True
            # output_text.write(line) -> for text file
            qn += 1
        elif tag.name == 'span': # if it's choice
            run = document.add_paragraph()
            run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = run.add_run(tag.text.strip())
            r.style = mystyle
            font = r.font
            font.rtl = True
            if cn >= len(CHOICE_LIST): # reset choice index
                cn = 0
            if tag.has_attr('style'):
                # \033[1mBOLD\033[0m -> BOLD ANSI code
                # choice = f'\033[1m{CHOICE_LIST[cn]}\033[0m' -> Trying to make a bold text -> failed in text format
                # choice = f'BOLD -> {CHOICE_LIST[cn]}' # -> for text file
                r = run.add_run(CHOICE_LIST[cn]).bold = True
                # r.bold = True
            else:
                # choice = CHOICE_LIST[cn] # -> for text file
                r = run.add_run(CHOICE_LIST[cn])
            
            # line = f'{choice}{tag.text.strip()}' # -> for text file
            # output_text.write(line) # -> for text file
            cn += 1

document.save('needtofix.docx')