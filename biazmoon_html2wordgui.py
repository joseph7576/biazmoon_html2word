import platform
import PySimpleGUI as psg
from docx import Document
from bs4 import BeautifulSoup
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

psg.theme('DarkBlue')
layout = [[psg.T("")],
          [psg.T("-"), psg.Text("Please choose your downloaded biazmoon's question html file: ")],
          [psg.T(""), psg.Input(), psg.FileBrowse(key="-htmlfile-", file_types=(('HTML Files', '*.html'),)), psg.Button("Convert")],
          [psg.T(""), psg.Quit("Exit")]]
window = psg.Window('Biazmoon html2word', layout, size=(550,150))


def open_file_explorer(directory):
    path = os.path.realpath(directory)
    os.startfile(path)

def create_word_doc(file, directory):
    document = Document()
    mystyle = document.styles.add_style('mystyle', WD_STYLE_TYPE.CHARACTER)
    with open(file, encoding='utf8') as fp: # you should replace your file with mine
        soup = BeautifulSoup(fp, 'html.parser')
        question_and_choices_list  = list(soup.find_all(['p', 'span']))

    CHOICE_LIST = [' (الف', ' (ب', ' (ج', ' (د'] # choice list for questions

    qn = 1 # question number
    cn = 0 # choice number(index)
    for tag in question_and_choices_list:
        if tag.name == 'p': # if it's question
            line = f'\n{qn}- {tag.text.strip()}'
            run = document.add_paragraph() # write to doc file
            run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = run.add_run(line)
            run.style = mystyle
            font = run.font
            font.rtl = True
            qn += 1
        elif tag.name == 'span': # if it's choice
            run = document.add_paragraph()
            run.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = run.add_run(tag.text.strip())
            r.style = mystyle
            font = r.font
            font.rtl = True
            if cn >= len(CHOICE_LIST): # reset choice index
                cn = 0
            if tag.has_attr('style'):
                r = run.add_run(CHOICE_LIST[cn]).bold = True
            else:
                r = run.add_run(CHOICE_LIST[cn])
            cn += 1
    
    sysinfo = platform.system()
    if sysinfo == "Linux":
        filepath = directory + '/needtofix.docx'
    elif sysinfo == "Windows":
        filepath  = directory + '\\needtofix.docx'
        
    document.save(filepath)


while True:
    event, values = window.read()
    if event == psg.WIN_CLOSED or event=="Exit" or event=="Quit":
        break
    elif event == "Convert":
        directory = psg.popup_get_folder("Choose a directory to save your docx file.", title='Where to save?')
        create_word_doc(values['-htmlfile-'], directory)
        psg.popup("needtofix.docx has created. Now go and fix that doc.", title=':D')
        open_file_explorer(directory)